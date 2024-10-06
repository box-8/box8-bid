from datetime import datetime
from io import BytesIO
import json
import os
from typing import List
from crewai import Agent, Crew, Process, Task
from docx import Document
import streamlit as st 
import tempfile
from crewai_tools import PDFSearchTool
from box8.Session import *
from box8.Agents import Commercial, Consultant, RagAgent
from box8.Functions import extraire_tableau_json, DocumentWriter

 
def do(uploaded_file):
    
        # Read the file and prepare the data for PDFSearchTool
        # PDF Search Tool setup
        doc = DocumentWriter()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
            f.write(uploaded_file.read())
        uploaded_file_path = f.name
        pdf_tool = PDFSearchTool(pdf=uploaded_file_path)
        
        # Création des agents
        agent_tech_genlecture = Agent(
            role="Expert Technique Ingéniérie",
            goal="""
            Synthétiser la fiche technique de l'équipement.""",
            verbose=True,
            memory=True,
            backstory="""
            Ingénieur en génie électromécanique et génie des procédés, vous avez une grande expertise dans l'analyse de 
            systèmes techniques complexes et faites appel à vos compétences de synthèse technique.
            """,
            max_iter=5,
            tools=[pdf_tool],
            llm = ChooseLLM()
        )
        agent_tech_genlecturetask = Task(
            description="Lire la documentation technique et synthétiser son contenu.",
            expected_output="""
            Un descriptif synthétique de la fiche d'équipements.""",
            agent=agent_tech_genlecture,
        )
        
        # Création des agents
        agent_tech_lecture = Agent(
            role="Expert Technique Ingéniérie",
            goal="""
            Identifier les composants techniques qui devront être inclus dans le processus de mise en service (commissionning) 
            d'aprés la fiche technique.
            Si la fiche technique n'indique pas de sous-équipements faites appel à vos connaissances 
            pour déterminer une décomposition en sous équipements.""",
            verbose=True,
            memory=True,
            backstory="""
            Ingénieur en génie électromécanique et génie des procédés, vous avez une grande expertise dans l'analyse de 
            systèmes techniques complexes et faites appel à vos connaissances pour déterminer 
            la composition d'un équipement en vue de sa mise en service (électricité, 
            instrumentation de mesure, instrumentation de régulation, organes mécaniques ...).
            """,
            max_iter=5,
            tools=[pdf_tool],
            llm = ChooseLLM()
        )
        agent_tech_lecturetask = Task(
            description="Lire la documentation technique et déterminer la composition de l'équipement.",
            expected_output="""
            Une liste avec l'intitulé du type d'équipement décrivant les composants et groupes 
            fonctionnels, organisée dans l'ordre de leur mise en service.""",
            agent=agent_tech_lecture,
        )
        
        
        crew = Crew(
            agents=[agent_tech_genlecture,agent_tech_lecture],
            tasks=[agent_tech_genlecturetask, agent_tech_lecturetask],
            process=Process.sequential  # Exécution séquentielle des tâches
        )
        
        part1 = "Groupes fonctionnels / équipements"
        st.subheader(part1)
        result = crew.kickoff()
        st.write(result.raw)
        
        doc.Chapter(part1)
        doc.writeBlack(result.raw)
        
        
        commissionning_agent = Agent(
            role="Ingénieur Mise en Service Commissionning",
            goal=f"""
            Rédiger une procédure de mise en service détaillée de l'équipement sachant qu'il comporte les éléments suivants :
            {result.raw} 
            Détailler des essais de contrôle de mise en service et de performance pour chacun d'eux
            """,
            verbose=True,
            memory=True,
            backstory="""
            Vous avez une grande expérience dans la mise en service de systèmes 
            industriels (chaleur, pétrochimie, pompage, énergie, transport, industrie manufacturière).
            """,
            max_iter=5,
            llm = ChooseLLM()
        )
        
        commissionning_agenttask = Task(
            description="Pour chaque composant de la liste décrire les tâches et actions à réaliser pour sa mise en service.",
            expected_output="""
            Une procédure de mise en service avec des instructions précises pour la mise en service. 
            """,
            agent=commissionning_agent,
            async_execution=False,  # S'exécute séquentiellement
        )
        crew = Crew(
            agents=[commissionning_agent],
            tasks=[commissionning_agenttask],
            process=Process.sequential  # Exécution séquentielle des tâches
        )
        
        part2 = "Mise en service / équipements"
        st.subheader(part2)
        result = crew.kickoff()
        st.write(result.raw)
        
        doc.Chapter(part2)
        doc.writeBlack(result.raw)
        
        
        redactor_agent = Agent(
            role="Technicien de Mise en Service",
            goal=f"""
            A partir de la liste ci aprés, décrire tâche par tâche le détail de la procédure de mise en service.
            Proposer des essais de performance pour chacun d'eux :
            {result.raw} 
            
            """,
            verbose=True,
            memory=True,
            backstory="""
            Vous avez une grande expérience dans la mise en service de systèmes 
            industriels (chaleur, pétrochimie, pompage, énergie, transport, industrie manufacturière).
            Vous entrez dans le détail des procédures de mise en service
            """,
            max_iter=5,
            llm = ChooseLLM()
        )
        
        redactor_agenttask = Task(
            description="""
            Pour chaque composant de la liste décrire les tâches et actions à réaliser 
            pour sa mise en service du point de vue technique, de la sécurité, de la qualité et des objectifs de performance.
            """,
            expected_output="""
            Une procédure technique de mise en service avec des instructions précises étape par étape. 
            """,
            agent=commissionning_agent,
            async_execution=False,  # S'exécute séquentiellement
        )
        # crew = Crew(
        #     agents=[agent_tech_lecture,redactor_agenttask],
        #     tasks=[agent_tech_lecturetask, redactor_agenttask],
        #     process=Process.sequential  # Exécution séquentielle des tâches
        # )
        crew = Crew(
            agents=[redactor_agent],
            tasks=[redactor_agenttask],
            process=Process.sequential  
        )
        part3 = "Détail de Mise en service"
        st.subheader(part3)
        result = crew.kickoff()
        st.write(result.raw)
        
        doc.Chapter(part3)
        doc.writeBlack(result.raw)
        
        
        if False:
            agent_verificateur = Agent(
                role="Vérificateur Qualité",
                goal=f"""
                S'assurer que la procédure est complète et conforme aux normes de rédaction technique : 
                {result.raw}
                """,
                verbose=True,
                memory=True,
                backstory="Spécialiste dans la validation de procédures techniques.",
                tools=[pdf_tool],
                llm = ChooseLLM()
                
            )
            agent_verificateurtask = Task(
                description="Vérifier la qualité de la procédure rédigée.",
                expected_output="Procédure de mise en service complète avec rapport de validation de conformité ou rapport de corrections.",
                agent=agent_verificateur,
            )
            
            crew = Crew(
                agents=[agent_verificateur],
                tasks=[agent_verificateurtask],
                process=Process.sequential  # Exécution séquentielle des tâches
            )
            
            part4 = "Vérification"
            st.subheader(part4)
            result = crew.kickoff()
            st.write(result.raw)
            doc.Chapter(part4)
            doc.writeBlack(result.raw)
        # doc.saveDocument(os.path.join("uploads","analyse-fiche-technique.docx"))
        return doc









def saveDoJson(data, name):
    # Créer un document Word
    doc = Document()
    
    # Ajouter le titre principal
    doc.add_heading(data['nom_equipement'], level=1)
    
    # Ajouter la synthèse
    doc.add_paragraph('Synthèse :')
    doc.add_paragraph(data['synthese'])

    # Ajouter les composants et leurs tâches de mise en service
    for composant in data['composants']:
        # Ajouter le nom du composant
        doc.add_heading(composant['nom_composant'], level=2)
        
        # Ajouter les tâches de mise en service, s'il y en a
        if composant['mise_en_service']:
            for tache in composant['mise_en_service']:
                # Ajouter le nom de la tâche
                doc.add_heading(tache["nom_tache"], level=3)
                # Ajouter la description de la tâche
                doc.add_paragraph(tache["description_tache"])
        else:
            # Si aucune mise en service n'est présente, indiquer "Aucune tâche de mise en service"
            doc.add_paragraph("Aucune tâche de mise en service disponible.")
    
    # Vérifier et créer le dossier "uploads" s'il n'existe pas
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    
    # Générer un nom de fichier unique avec un timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    file_path = os.path.join("uploads", f"{name}-mes-{timestamp}.docx")

    # Sauvegarder le document
    doc.save(file_path)
    
    print(f"Le document a été sauvegardé sous : {file_path}")








def doJson(backstory, uploaded_file):

    # Read the file and prepare the data for PDFSearchTool
    # PDF Search Tool setup
    # doc = DocumentWriter()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        f.write(uploaded_file.read())
    uploaded_file_path = f.name
    pdf_tool = PDFSearchTool(pdf=uploaded_file_path)
    
    # Création des agents
    agent_tech_genlecture = Agent(
        role="Expert Technique Ingéniérie",
        goal="""
        Synthétiser la fiche technique de l'équipement.""",
        verbose=True,
        memory=True,
        backstory="""
        Ingénieur en génie électromécanique et génie des procédés, vous avez une grande expertise dans l'analyse de 
        systèmes techniques complexes et faites appel à vos compétences pour en faire la synthèse technique.
        """,
        max_iter=5,
        tools=[pdf_tool],
        llm = ChooseLLM()
    )
    agent_tech_genlecturetask = Task(
        description="Lire la documentation technique et synthétiser son contenu.",
        expected_output="""
        Un descriptif synthétique de la fiche d'équipements.""",
        agent=agent_tech_genlecture,
    )
    
    # Création des agents
    agent_tech_lecture = Agent(
        role="Expert Technique Ingéniérie",
        goal="""
        Identifier les composants techniques qui devront être inclus dans le processus de mise en service (commissionning) 
        d'aprés la fiche technique.
        Si la fiche technique n'indique pas de sous-équipements faites appel à vos connaissances 
        pour déterminer une décomposition en sous équipements.""",
        verbose=True,
        memory=True,
        backstory=backstory,
        max_iter=5,
        tools=[pdf_tool],
        llm = ChooseLLM()
    )
    agent_tech_lecturetask = Task(
        description="Lire la documentation technique et déterminer la composition de l'équipement.",
        expected_output="""
        Un objet JSON bien formaté avec l'intitulé du type d'équipement décrivant les composants et groupes 
        fonctionnels, organisée dans l'ordre de leur mise en service. Le format de retour est le suivant : 
        {{
            "nom_equipement":"nom de l'équipement principal objet la fiche technique",
            "synthese":"descriptif synthetique de la nature et de la fonction de l'équipement",
            "composants":[
                {{
                    "nom_composant":"nom du composant dans l'équipement principal",
                    "mise_en_service":[tableau vierge]
                }}
                ...
            ] 
        }}""",
        agent=agent_tech_lecture,
    )
    crew = Crew(
        agents=[agent_tech_genlecture,agent_tech_lecture],
        tasks=[agent_tech_genlecturetask, agent_tech_lecturetask],
        process=Process.sequential  # Exécution séquentielle des tâches
    )
    
    part1 = "Groupes fonctionnels / équipements"
    st.subheader(part1)
    result = crew.kickoff()

    data = json.loads(result.raw)
    
    st.success("Procédure incomplète")
    st.write(data)
    # st.write(f"analyse {eq}")
    # st.write(eqSyn)

    
    st.subheader("Procédure détaillée")
    
    backstory=f"""
        Vous avez une grande expérience dans la mise en service de systèmes 
        industriels (chaleur, pétrochimie, pompage, énergie, transport, industrie manufacturière).
        Vous mettez en service un {data["nom_equipement"]}. {data["synthese"]}
        """
    st.success(backstory)
    if True:
        for composant in data["composants"]:
            # element = composant["nom_composant"]
            # st.write(element)
            commission(backstory, composant)
    
    return data, f.name
    
    
def commission(backstory, composant):
    
    goal=f"""
        Détailler des essais de mise en service et tests de performance pour {composant["nom_composant"]}
        """
    
    commissionning_agent = Agent(
        role="Ingénieur Mise en Service",
        goal=goal,
        verbose=True,
        memory=True,
        backstory=backstory,
        max_iter=5,
        llm = ChooseLLM()
    )
    
    commissionning_agenttask = Task(
        description="Pour chaque composant de l'équipements décrire les tâches et actions à réaliser pour sa mise en service.",
        expected_output="""
        Un tableau JSON avec les tâches de mise en service au format suivant : 
        [
            {{
                "nom_tache":"nom de la tâche",
                "description_tache": "description détaillée de la tâche de mise en service du composant"
            }}
            ...
        ] 
        """,
        agent=commissionning_agent,
        async_execution=False,  # S'exécute séquentiellement
    )
    crew = Crew(
        agents=[commissionning_agent],
        tasks=[commissionning_agenttask],
        process=Process.sequential  # Exécution séquentielle des tâches
    )
    result = crew.kickoff()
    datatab = extraire_tableau_json(result.raw)
    composant["mise_en_service"] = datatab
    # ici
    st.warning(goal) 
    st.write(datatab)
    return composant
    
    


st.set_page_config(page_title="Analyse de fiche technique", page_icon="🤖", layout="wide") 
init_session()
st.header("Mise en service")

backstory_default="""Ingénieur en génie électromécanique et génie des procédés, vous avez une grande expertise dans l'analyse de 
systèmes techniques complexes et faites appel à vos connaissances pour déterminer  
la composition d'un équipement en vue de sa mise en service (électricité, 
instrumentation de mesure, instrumentation de régulation, organes mécaniques ...)."""

backstory = st.text_area("but de l'analyse",backstory_default)


uploaded_file_mes = st.file_uploader("Upload technical document", type="pdf")
if st.button("Analyse de la fiche technique"):
    
    # Streamlit file uploader for PDF
    
    if uploaded_file_mes :
        
        data, name = doJson(backstory, uploaded_file_mes)
        saveDoJson(data, uploaded_file_mes.name)
        
        # doc.saveDocument(os.path.join("uploads","analyse-fiche-technique.docx"))
    
