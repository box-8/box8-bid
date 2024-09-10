import streamlit as st
from crewai import Agent
from crewai_tools import PDFSearchTool
from io import BytesIO
from docx import Document

# --- Streamlit Interface ---
st.title("Simulation de Consultation Médicale")

# Upload patient medical data PDF
patient_pdf = st.file_uploader("Charger le PDF du Patient (Antécédents Médicaux, Résultats d'examen...)", type="pdf")

# Upload doctor specialization PDF
doctor_pdf = st.file_uploader("Charger le PDF du Médecin (Spécialité médicale)", type="pdf")

# Button to trigger analysis
if st.button("Analyser les données du patient"):
    if not patient_pdf or not doctor_pdf:
        st.error("Veuillez télécharger les documents du patient et du médecin.")
    else:
        # Agents setup
        # Simulate Patient agent with their medical documents
        patient_agent = Agent(
            role="Patient",
            goal="Fournir un contexte médical basé sur les antécédents médicaux",
            tools=[PDFSearchTool(pdf=patient_pdf)]
        )
        
        # Simulate Doctor agent with specialization in analyzing patient data
        doctor_agent = Agent(
            role="Médecin",
            goal="Analyser les données médicales et fournir un diagnostic",
            tools=[PDFSearchTool(pdf=doctor_pdf)]
        )

        # Crew interaction (Patient data passed to Doctor)
        patient_data = patient_agent.execute_task("Rechercher les informations médicales importantes")
        doctor_response = doctor_agent.execute_task(f"Analyse des données du patient: {patient_data}")
        
        # Display doctor's conclusions
        st.subheader("Diagnostic et Plan de Traitement du Médecin")
        st.write(doctor_response)

        # Document generation (medical reasoning)
        doc = Document()
        doc.add_heading('Rapport Médical', 0)
        doc.add_paragraph(f'Diagnostic et Plan de Traitement:\n{doctor_response}')
        doc.add_paragraph('Explication du Raisonnement Médical:')
        doc.add_paragraph("Le médecin a analysé les données en se basant sur les antécédents médicaux et les résultats d'examen fournis dans le PDF.")
        
        # Downloadable medical report (Word)
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        st.download_button("Télécharger le rapport médical", data=doc_stream, file_name="rapport_medical.docx")

