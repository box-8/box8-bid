
from datetime import datetime
import json
import streamlit as st
from docx import Document
from docx.shared import RGBColor
from crewai_tools import PDFSearchTool


if 'llm_model' not in st.session_state:
    st.session_state.llm_model = 'openai'  
    
Toaster: str = None  # Déclaration de la variable globale avec typage explicite
def toast(val=None):
    global Toaster  # Utilisation du mot-clé global pour accéder à la variable globale
    if val is None:
        st.toast(Toaster)
    else:
        Toaster += "\n " + val  # Concaténation améliorée 
    



def extraire_tableau_json(chaine):
    """
    Extrait le premier tableau JSON trouvé dans une chaîne de caractères.
    Args:
        chaine: La chaîne de caractères contenant le tableau JSON.

    Returns:
        Le tableau JSON extrait sous forme de liste Python, ou None si aucun tableau n'est trouvé.
    """
    debut = chaine.find('[')
    if debut == -1:
        return None  # Aucun tableau trouvé

    fin = chaine.rfind(']')
    if fin == -1 or fin < debut:
        return None  # Tableau mal formé
    try:
        tableau_json = json.loads(chaine[debut:fin+1])
        return tableau_json
    except json.JSONDecodeError:
        return None  # Erreur lors de la désérialisation JSON



class DocumentWriter():
    def __init__(self, title=""):
        self.docAnalyse = Document()
        self.Title(title)

    def Title(self,title):
        self.docAnalyse.add_heading(title, 0)
        
    def Chapter(self,title):
        self.docAnalyse.add_heading(title, 1)
    
    def SubChapter(self,title):
        self.docAnalyse.add_heading(title, 2)
    
    def writeBlack(self,text):
        paragraph = self.docAnalyse.add_paragraph(text)
        paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    
    def writeBlue(self,text):
        paragraph = self.docAnalyse.add_paragraph(text)
        paragraph.style.font.color.rgb = RGBColor(0, 0, 75)
        
    def writeDocument(self, text, heading_level=None, color=RGBColor(0, 0, 0)): # Noir par défaut
        """
        Ajoute du texte au document global avec une mise en forme spécifique.

        Args:
            text (str): Le texte à ajouter au document.
            heading_level (int, optional): Le niveau de titre (1 à 9). Par défaut, un titre de niveau 1.
            color (RGBColor, optional): La couleur du texte. Par défaut, rouge vif.
        """
        

        if heading_level is None:
            paragraph = self.docAnalyse.add_paragraph(text)
            paragraph.style.font.color.rgb = color
        elif heading_level == 0:
            self.docAnalyse.add_heading(text, 0)
        elif 1 <= heading_level <= 3:
            heading = self.docAnalyse.add_heading(text, heading_level)
            heading.style.font.color.rgb = color
        else:
            paragraph = self.docAnalyse.add_paragraph(text)
            paragraph.style.font.color.rgb = color

    # Exemple d'utilisation
    # writeDocument("Ceci est un titre important", heading_level=1)
    # writeDocument("Et voici un paragraphe en rouge vif.")


    def saveDocument(self,name=None):
        self.docAnalyse  # Indique que vous utilisez la variable globale à l'intérieur de la fonction
        # Save the Word document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"result_{timestamp}.docx"
        
        if name is not None : 
            path = self.docAnalyse.save(f"{name}-{timestamp}.docx")
        else:
            path = self.docAnalyse.save(filename)
        self.path = path
        return path

