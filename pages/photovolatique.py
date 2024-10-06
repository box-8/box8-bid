from datetime import datetime
from io import BytesIO
import json
import os
from typing import List
from crewai import Agent, Crew, Process, Task
from docx import Document
import streamlit as st 
import tempfile
from crewai_tools import PDFSearchTool
from box8.Session import *
from box8.Agents import Commercial, Consultant, RagAgent
from box8.Functions import extraire_tableau_json, DocumentWriter






def saveDoJson(data, name):
    # Créer un document Word
    doc = Document()
    
    # Ajouter le titre principal
    doc.add_heading(data['nom_equipement'], level=1)
    
    # Ajouter la synthèse
    doc.add_paragraph('Synthèse :')
    doc.add_paragraph(data['synthese'])

    # Ajouter les composants et leurs tâches de mise en service
    for composant in data['composants']:
        # Ajouter le nom du composant
        doc.add_heading(composant['nom_composant'], level=2)
        
        # Ajouter les tâches de mise en service, s'il y en a
        if composant['mise_en_service']:
            for tache in composant['mise_en_service']:
                # Ajouter le nom de la tâche
                doc.add_heading(tache["nom_tache"], level=3)
                # Ajouter la description de la tâche
                doc.add_paragraph(tache["description_tache"])
        else:
            # Si aucune mise en service n'est présente, indiquer "Aucune tâche de mise en service"
            doc.add_paragraph("Aucune tâche de mise en service disponible.")
    
    # Vérifier et créer le dossier "uploads" s'il n'existe pas
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    
    # Générer un nom de fichier unique avec un timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    file_path = os.path.join("uploads", f"{name}-mes-{timestamp}.docx")

    # Sauvegarder le document
    doc.save(file_path)
    
    print(f"Le document a été sauvegardé sous : {file_path}")








def doJson(backstory, uploaded_file):

    # Read the file and prepare the data for PDFSearchTool
    # PDF Search Tool setup
    # doc = DocumentWriter()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        f.write(uploaded_file.read())
    uploaded_file_path = f.name
    pdf_tool = PDFSearchTool(pdf=uploaded_file_path)
    
    # Création des agents
    agent_tech_genlecture = Agent(
        role="Ingénieur Développement projets photovoltaique",
        goal="""
        Synthétiser l'appel d'ofre pour le champs photovolatique pour permettre son chiffrage.""",
        verbose=True,
        memory=True,
        backstory="""
        Ingénieur en génie électrique et énergies renouvellables, vous avez une grande expertise dans le montage de dossiers réponse 
        aux d'appels d'offre de la Commission de Régulation de l'Energie 
        """,
        max_iter=5,
        tools=[pdf_tool],
        llm = ChooseLLM()
    )
    agent_tech_genlecturetask = Task(
        description="Lire l'appel d'offre et synthétiser son contenu pour le bureau d'études.",
        expected_output="""
        Un objet JSON bien formaté contenant les informations extraites du document : 
        {{
            "objet_appel_offre":"contexte du dossier d'appel d'offre",
            "Examen_des_offres":"résumé de la procédure d'examen",
            "Admissibilité":"résumé des condictions d'admissibilité des candidats",
            "Notation":"résumé des critères de notation des offres",
            "Rémunération":"résumé de la rémunération",
            "taches":[tableau vide]
        }}
        """,
        agent=agent_tech_genlecture,
    )
    
    # Création des agents
    agent_tech_lecture = Agent(
        role="Ingénieur Etudes",
        goal="""
        Identifier les tâches qui devront être incluses dans le processus de construction (études complémentaires, plans, terrassements, fondations, structures, réseaux enterrés, panneaux, câblage, boitiers de raccordement, 
onduleurs, postes transformation, instrumentation de mesure, instrumentation de régulation, clôtures...).""",
        verbose=True,
        memory=True,
        backstory=backstory,
        max_iter=5,
        tools=[pdf_tool],
        llm = ChooseLLM()
    )
    agent_tech_lecturetask = Task(
        description="Ajouter une liste des taches au format JSON",
        expected_output="""
        Un objet JSON bien formaté avec l'intitulé du type d'équipement décrivant les composants et groupes 
        fonctionnels, organisée dans l'ordre de leur mise en service. Le format de retour est le suivant : 
        [
                {{
                    "construction_task":"tâche de construction (études complémentaires, plans, terrassements, fondations, structures, réseaux enterrés, panneaux, câblage, boitiers de raccordement, 
onduleurs, postes transformation, instrumentation de mesure, instrumentation de régulation, clôtures)",
                    "prix":"estimation de prix pour la tâche"
                }}
                ...
        ] """,
        agent=agent_tech_lecture,
    )
    crew = Crew(
        agents=[agent_tech_genlecture,agent_tech_lecture],
        tasks=[agent_tech_genlecturetask, agent_tech_lecturetask],
        process=Process.sequential  # Exécution séquentielle des tâches
    )   
    
    part1 = "Etude de l'AO CRE"
    st.subheader(part1)
    result = crew.kickoff()

    data = json.loads(result.raw)
    
    st.write(data)
    # st.write(f"analyse {eq}")
    return data, f.name
    
    st.subheader("Procédure détaillée")
    
    backstory=f"""
        Vous avez une grande expérience dans la mise en service de systèmes 
        industriels (chaleur, pétrochimie, pompage, énergie, transport, industrie manufacturière).
        Vous mettez en service un {data["nom_equipement"]}. {data["synthese"]}
        """
    st.success(backstory)
    if True:
        for composant in data["composants"]:
            # element = composant["nom_composant"]
            # st.write(element)
            commission(backstory, composant)
    
    return data, f.name
    


st.set_page_config(page_title="Analyse de fiche technique", page_icon="🤖", layout="wide") 
init_session()
st.header("Photovoltaique")

backstory_default="""Ingénieur développement de projets photovoltaiques vous avez une grande expertise dans l'analyse des cahier des charges CRE 
et faites appel à votre bureau d'études pour pour chiffrer la construction des champs PV de l'appel d'offre."""

backstory = st.text_area("but de l'analyse",backstory_default)


uploaded_file_mes = st.file_uploader("Upload technical document", type="pdf")
if st.button("Analyse de la fiche technique"):
    
    # Streamlit file uploader for PDF
    
    if uploaded_file_mes :
        
        data, name = doJson(backstory, uploaded_file_mes)
        # Chemin relatif vers le répertoire 'uploads'
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        path = os.path.join('uploads', f'fichier-cre-{timestamp}.json')

        # Vérifier si le répertoire 'uploads' existe, sinon le créer
        if not os.path.exists('uploads'):
            os.makedirs('uploads')

        # Sauvegarder les données dans un fichier JSON dans le répertoire 'uploads'
        with open(path, 'w') as fichier:
            json.dump(data, fichier, indent=4, ensure_ascii=False)
    
